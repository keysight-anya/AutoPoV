"""
AutoPoV Documentation Generator
Produces both PDF (via fpdf2) and DOCX (via python-docx)
Output: DOCS_APPLICATION_FLOW.pdf  and  DOCS_APPLICATION_FLOW.docx
"""

# ============================================================
# CONTENT  (shared by both outputs)
# ============================================================
# Each item is  (type, payload)
# Types:
#   cover        – (title, subtitle, sub2)
#   h1           – section heading
#   h2           – sub-heading
#   body         – paragraph text
#   bullets      – list[str]
#   table        – {"headers": [...], "rows": [[...]], "col_widths_pdf": [...], "col_widths_docx_cm": [...]}
#   code_block   – list[str]  (monospace block)
#   pipeline     – special ASCII pipeline diagram block
#   page_break   – force new page (PDF only; DOCX always honours natural flow)

SECTIONS = [

    ("cover", ("AutoPoV",
               "Automated Proof-of-Vulnerability Framework",
               "Application Flow and System Reference")),

    # ------------------------------------------------------------------
    ("h1", "1.  System Overview"),
    ("body",
     "AutoPoV is a multi-agent vulnerability scanner. It accepts a codebase, discovers "
     "security weaknesses, confirms whether each one is a real exploitable issue, and then "
     "generates a working proof that the vulnerability can be triggered. Once a scan is "
     "submitted, the entire process runs without further input."),
    ("body",
     "The system is built on a LangGraph state machine. The graph is compiled once when "
     "the server starts and is reused for every scan. Each node in the graph reads the "
     "current scan state, performs its work, writes its results back into state, and "
     "passes control to the next node via a routing condition."),
    ("body", "The core agent loop:"),
    ("bullets", [
        "Read the current scan state",
        "Evaluate the routing condition to decide the next step",
        "Execute the work (tool call or language model query)",
        "Write results back into state",
        "Route to the next node and repeat",
    ]),

    # ------------------------------------------------------------------
    ("h1", "2.  Stage 1 -- Scan Intake"),
    ("body",
     "A scan starts when code is submitted through one of four entry points. The API layer "
     "validates the request, creates a scan record, resolves the code to a local path, "
     "and launches the agent graph in a background thread. The scan ID is returned to the "
     "caller immediately so they do not have to wait for analysis to complete."),
    ("table", {
        "headers": ["Entry Point", "API Endpoint", "Source"],
        "rows": [
            ["Git repository",  "POST /api/scan/git",                 "Cloned from a URL"],
            ["ZIP archive",     "POST /api/scan/zip",                 "Uploaded file"],
            ["Raw code paste",  "POST /api/scan/paste",               "Inline text string"],
            ["Webhook event",   "POST /api/webhook/github  or  /gitlab", "Push or pull-request event"],
        ],
        "col_widths_pdf":     [110, 180, 150],
        "col_widths_docx_cm": [4.0, 6.5,  5.5],
    }),
    ("body", "On every submission the API layer:"),
    ("bullets", [
        "Validates the API key using a SHA-256 hash comparison",
        "Enforces the per-key rate limit (10 scans per 60 seconds)",
        "Creates a scan record with a unique UUID and initial status 'created'",
        "Clones the repo, extracts the ZIP, or saves the pasted code to a workspace path on disk",
        "Dispatches the agent graph as a background task",
        "Returns {\"scan_id\": \"...\"} immediately",
    ]),

    # ------------------------------------------------------------------
    ("h1", "3.  Stage 2 -- Code Ingestion"),
    ("body",
     "Before any analysis begins, the Ingestion Agent prepares the codebase for semantic "
     "search. This preparation allows every later agent to retrieve the precise code context "
     "it needs without having to read entire files on demand."),
    ("body", "The Ingestion Agent:"),
    ("bullets", [
        "Walks every file in the codebase, skipping binaries, build folders, and lock files",
        "Splits each file into overlapping chunks of 4,000 characters with a 200-character overlap, "
        "splitting on class and function boundaries first",
        "Embeds each chunk using openai/text-embedding-3-small (online mode) or "
        "sentence-transformers/all-MiniLM-L6-v2 (offline mode)",
        "Stores the chunks and their vectors in ChromaDB under a collection scoped to the scan ID",
    ]),
    ("body",
     "If embedding fails for any reason, the scan continues. All later agents fall back to "
     "reading source files directly from disk."),

    # ------------------------------------------------------------------
    ("h1", "4.  Stage 3 -- Vulnerability Discovery"),
    ("body",
     "Three independent discovery strategies run in sequence. Their findings are merged into "
     "a single list and any duplicates sharing the same file path, line number, and CWE type "
     "are removed before investigation begins."),

    ("h2", "4a.  CodeQL Static Discovery"),
    ("body",
     "CodeQL builds a database from the source code and runs structured queries against it. "
     "Because it understands the code's syntax tree and data flow, it can trace a value from "
     "user input through a call chain to a dangerous sink. This makes it the highest-precision "
     "discovery method in the system."),
    ("bullets", [
        "Detect the dominant programming language from file extensions",
        "Build a CodeQL database from the source root",
        "Run one language-specific query per requested CWE",
        "Parse the SARIF output; each result becomes a candidate finding with confidence 0.8",
        "Remove the database once all queries have completed",
    ]),
    ("body",
     "If CodeQL is not installed, this step is skipped and the system relies on the two "
     "methods below."),

    ("h2", "4b.  Heuristic Scout"),
    ("body",
     "The Heuristic Scout applies regex pattern libraries against every source file. "
     "It covers all 20 supported CWEs, runs in milliseconds, and has no cost. "
     "It catches patterns that CodeQL can miss, such as Python-specific string concatenation "
     "in SQL queries. Candidates produced here have a confidence of 0.35 and must be confirmed "
     "by the Investigator Agent before proceeding."),

    ("h2", "4c.  LLM Scout"),
    ("body",
     "If enabled, the LLM Scout sends batches of source file snippets to the language model "
     "and asks it to identify potential vulnerabilities. It acts as a last resort when neither "
     "CodeQL nor heuristics find anything. The scout aborts automatically if the accumulated "
     "cost would exceed the configured SCOUT_MAX_COST_USD limit."),
    ("body", "Supported CWEs:"),
    ("bullets", [
        "CWE-89  SQL Injection            CWE-79  Cross-Site Scripting (XSS)",
        "CWE-22  Path Traversal           CWE-78  Command Injection",
        "CWE-94  Code Injection           CWE-502 Unsafe Deserialization",
        "CWE-798 Hardcoded Credentials    CWE-312 Cleartext Storage",
        "CWE-327 Weak Cryptography        CWE-352 CSRF",
        "CWE-287 Authentication Failure   CWE-306 Missing Authentication",
        "CWE-601 Open Redirect            CWE-918 Server-Side Request Forgery",
        "CWE-434 Unrestricted File Upload CWE-611 XML External Entity (XXE)",
        "CWE-400 Resource Exhaustion      CWE-384 Session Fixation",
        "CWE-200 Information Disclosure   CWE-20  Input Validation",
    ]),

    # ------------------------------------------------------------------
    ("h1", "5.  Stage 4 -- Investigation"),
    ("body",
     "The Investigator Agent reviews each candidate finding one at a time. For each one, it "
     "assembles all available context and asks a language model to decide: is this a real "
     "vulnerability or a false positive? The agent does not pass raw code to the model. It "
     "builds a structured prompt that includes the flagged code, surrounding lines, and "
     "semantically related code retrieved from ChromaDB."),
    ("bullets", [
        "Retrieve the 50 lines surrounding the flagged line from the source file on disk",
        "Query ChromaDB for semantically related code chunks from the same codebase",
        "For CWE-416 (Use-After-Free) only: run Joern to trace data flow through the call graph",
        "Send all assembled context to the language model with a structured investigation prompt",
        "Parse the model response for verdict (REAL or FALSE_POSITIVE), confidence (0 to 1), "
        "explanation, root cause, and the exact vulnerable code snippet",
        "Record the model name, cost, verdict, and confidence in the Learning Store",
    ]),
    ("body",
     "Only findings where the verdict is REAL and confidence is 0.7 or higher are passed to "
     "the PoV Generator Agent. All others are marked as skipped."),

    ("h2", "5a.  Model Selection -- Policy Agent"),
    ("body",
     "Before each investigation call, the Policy Agent selects which language model to use. "
     "In learning mode, it queries the Learning Store for the model that has produced the "
     "best ratio of confirmed vulnerabilities to cost for the current CWE and programming "
     "language. The more scans the system has processed, the more informed this decision becomes."),
    ("table", {
        "headers": ["Mode", "Behaviour"],
        "rows": [
            ["auto",     "OpenRouter selects the best available model automatically"],
            ["fixed",    "Always uses the model configured in MODEL_NAME"],
            ["learning", "Queries the Learning Store and selects the model with the best "
                         "confirmed-vulnerability-to-cost score for this CWE and language"],
        ],
        "col_widths_pdf":     [70, 360],
        "col_widths_docx_cm": [2.5, 13.5],
    }),

    # ------------------------------------------------------------------
    ("h1", "6.  Stage 5 -- Proof-of-Vulnerability Generation"),
    ("body",
     "For each confirmed finding, the PoV Generator Agent writes a Python script that "
     "actually triggers the vulnerability. The purpose is not only to identify the problem "
     "but to demonstrate that it can be reliably exploited. The script must print "
     "VULNERABILITY TRIGGERED to standard output when it succeeds."),
    ("bullets", [
        "Retrieve the full source file from ChromaDB or disk",
        "Detect the target codebase language",
        "Send the vulnerable code, the Investigator's explanation, and language context to the model",
        "Strip any markdown formatting from the response",
        "Store the clean script in the finding state",
    ]),
    ("body",
     "PoV scripts are restricted to the Python standard library. This constraint ensures "
     "the script can execute in any sandboxed environment without requiring a package install step."),
    ("body",
     "If generation fails, the finding is marked pov_generation_failed and the scan "
     "continues with the remaining findings."),

    # ------------------------------------------------------------------
    ("h1", "7.  Stage 6 -- Validation"),
    ("body",
     "The Validation Agent checks that the generated PoV script is correct and actually "
     "triggers the vulnerability. It applies three methods in order, escalating only if "
     "the previous method does not produce a confident result."),
    ("table", {
        "headers": ["Tier", "Method", "Confirmation condition"],
        "rows": [
            ["1  Static",
             "Checks the script for correct exploit patterns, the required print statement, "
             "stdlib-only imports, and CWE-specific structural patterns",
             "Static confidence >= 80%"],
            ["2  Unit Test",
             "Injects the vulnerable code into a subprocess namespace and executes the PoV script",
             "VULNERABILITY TRIGGERED appears in stdout"],
            ["3  Docker",
             "Runs the PoV inside a sandboxed container with no network access, "
             "512 MB RAM, 1 CPU, and a 60-second timeout",
             "VULNERABILITY TRIGGERED appears in container stdout"],
        ],
        "col_widths_pdf":     [60, 230, 150],
        "col_widths_docx_cm": [2.2, 8.5,  5.3],
    }),
    ("body",
     "If validation fails after the third tier, the PoV Generator Agent is called again "
     "with feedback about what went wrong. This retry loop runs up to two times before "
     "the finding is marked as failed."),

    # ------------------------------------------------------------------
    ("h1", "8.  Stage 7 -- Loop Routing"),
    ("body",
     "After each finding is resolved (confirmed, skipped, or failed), the graph's conditional "
     "router increments the finding index and checks whether more findings remain. If they do, "
     "it routes back to the Investigator Agent for the next one. When all findings have been "
     "processed, the scan is marked completed and the graph exits."),
    ("body",
     "This loop is what separates the system from a linear script. The graph cycles through "
     "every finding, running the full investigation and validation pipeline on each one, "
     "and can route differently depending on what each finding reveals."),

    # ------------------------------------------------------------------
    ("h1", "9.  Stage 8 -- Results and Reports"),
    ("body",
     "When the agent graph finishes, the Scan Manager persists the full results to disk and "
     "makes them available via the API."),
    ("bullets", [
        "Full result JSON written to  results/runs/<scan_id>.json",
        "Summary row appended to  results/runs/scan_history.csv",
        "Confirmed PoV scripts saved to  results/povs/",
        "Optional codebase snapshot at  results/snapshots/<scan_id>/  "
        "(required for scan replay; enabled by SAVE_CODEBASE_SNAPSHOT=true)",
    ]),
    ("body",
     "Reports can be downloaded at any time after the scan completes via "
     "GET /api/report/<scan_id>?format=json  or  ?format=pdf. "
     "The PDF report includes an executive summary, all confirmed vulnerabilities with their "
     "PoV scripts, false positive analysis, per-model cost breakdown, and scan metadata."),

    # ------------------------------------------------------------------
    ("h1", "10.  Real-Time Observability"),
    ("body",
     "Every log message produced by an agent is written to three places simultaneously: "
     "the LangGraph state's log list, the Scan Manager's thread-safe in-memory buffer, "
     "and an SSE stream served at GET /api/scan/<scan_id>/stream. Both the web dashboard "
     "and the CLI subscribe to this stream and display agent activity as the scan runs. "
     "No polling is required; events are pushed as they occur."),

    # ------------------------------------------------------------------
    ("h1", "11.  Self-Improvement via the Learning Store"),
    ("body",
     "Every investigation result and PoV execution outcome is recorded in a SQLite database "
     "at data/learning.db. The Policy Agent queries this database when it needs to select a "
     "model. In learning routing mode, it picks the model that has delivered the best ratio "
     "of confirmed vulnerabilities to cost for the specific CWE type and programming language "
     "being scanned."),
    ("code_block", [
        "SELECT model, confirmed / (cost + 0.01) AS score",
        "FROM   investigations",
        "WHERE  cwe = ? AND language = ?",
        "GROUP  BY model",
        "ORDER  BY score DESC",
        "LIMIT  1",
    ]),
    ("body",
     "Because every scan adds new records, the routing decisions improve over time. "
     "A system that has scanned many Python codebases for SQL injection will reliably "
     "choose the model that has worked best in that specific context."),

    # ------------------------------------------------------------------
    ("h1", "12.  End-to-End Pipeline"),
    ("body",
     "The diagram below shows how data flows through the system from submission to report."),
    ("pipeline", [
        "Code Submitted",
        "  (Git URL / ZIP Upload / Code Paste / Webhook)",
        "",
        "         |",
        "         v",
        "",
        "+------------------+       +------------------------------+",
        "|  Ingestion Agent | ----> |  ChromaDB  (per-scan store)  |",
        "+------------------+       +------------------------------+",
        "",
        "         |",
        "         v",
        "",
        "+------------------------------------------------+",
        "|  Discovery                                     |",
        "|    CodeQL       - AST + dataflow, SARIF output |",
        "|    Heuristic    - regex patterns, 20 CWEs      |",
        "|    LLM Scout    - language model (fallback)    |",
        "+------------------------------------------------+",
        "         |  merged + deduplicated",
        "         v",
        "",
        "+-------------------+   <---  ChromaDB context",
        "|  Investigator     |   <---  Policy Agent (model selection)",
        "|  Agent            |",
        "+-------------------+",
        "         |  verdict: REAL, confidence >= 0.7",
        "         v",
        "",
        "+---------------------+",
        "|  PoV Generator      |",
        "|  Agent              |",
        "+---------------------+",
        "         |",
        "         v",
        "",
        "+---------------------------------------------+",
        "|  Validation Agent                           |",
        "|    Tier 1: Static analysis                  |",
        "|    Tier 2: Unit test (subprocess)           |",
        "|    Tier 3: Docker container (sandboxed)     |",
        "+---------------------------------------------+",
        "         |",
        "         +---> more findings? ---> back to Investigator",
        "         |",
        "         v  all findings processed",
        "",
        "+------------------+       +---------------------------+",
        "|  Learning Store  | <---  |  Records every outcome    |",
        "+------------------+       +---------------------------+",
        "",
        "         |",
        "         v",
        "",
        "+--------------------+",
        "|  Report Generator  | ---->  JSON  +  PDF  download",
        "+--------------------+",
    ]),

    # ------------------------------------------------------------------
    ("h1", "13.  API Reference"),
    ("table", {
        "headers": ["Method", "Endpoint", "Auth", "Description"],
        "rows": [
            ["POST",   "/api/scan/git",          "API Key",    "Scan a Git repository"],
            ["POST",   "/api/scan/zip",           "API Key",    "Scan a ZIP archive"],
            ["POST",   "/api/scan/paste",         "API Key",    "Scan pasted code"],
            ["GET",    "/api/scan/{id}",          "API Key",    "Get scan status and findings"],
            ["GET",    "/api/scan/{id}/stream",   "API Key *",  "Stream live agent logs (SSE)"],
            ["POST",   "/api/scan/{id}/cancel",   "API Key",    "Cancel a running scan"],
            ["POST",   "/api/scan/{id}/replay",   "API Key",    "Re-run findings on new models"],
            ["GET",    "/api/history",            "API Key",    "Paginated scan history"],
            ["GET",    "/api/report/{id}",        "API Key",    "Download report (JSON or PDF)"],
            ["GET",    "/api/learning/summary",   "API Key",    "Model performance statistics"],
            ["GET",    "/api/metrics",            "API Key",    "System-wide scan metrics"],
            ["GET",    "/api/config",             "API Key",    "Config and tool availability"],
            ["GET",    "/api/health",             "None",       "Server health check"],
            ["POST",   "/api/keys/generate",      "Admin Key",  "Create a new API key"],
            ["GET",    "/api/keys",               "Admin Key",  "List all API keys"],
            ["DELETE", "/api/keys/{id}",          "Admin Key",  "Revoke an API key"],
            ["POST",   "/api/admin/cleanup",      "Admin Key",  "Remove old result files"],
            ["POST",   "/api/webhook/github",     "HMAC",       "GitHub push triggers a scan"],
            ["POST",   "/api/webhook/gitlab",     "Token",      "GitLab push triggers a scan"],
        ],
        "col_widths_pdf":     [44, 155, 68, 173],
        "col_widths_docx_cm": [1.6, 5.6, 2.5, 6.3],
    }),
    ("body",
     "* The SSE stream endpoint accepts the API key as a query parameter: "
     "?api_key=<key>  because browser EventSource does not support custom headers."),
    ("body", "Interactive API documentation: http://localhost:8000/api/docs"),

    # ------------------------------------------------------------------
    ("h1", "14.  Authentication"),
    ("table", {
        "headers": ["Key Type", "Stored As", "Rate Limit", "Used For"],
        "rows": [
            ["Admin Key",
             "Plaintext in .env",
             "No limit",
             "Key management endpoints only"],
            ["API Key",
             "SHA-256 hash in data/api_keys.json",
             "10 scans per 60 seconds per key",
             "All scan, report, and data endpoints"],
        ],
        "col_widths_pdf":     [75, 140, 115, 110],
        "col_widths_docx_cm": [2.8, 5.2, 4.2, 3.8],
    }),
    ("body",
     "API keys are never stored in plain text. The server hashes the submitted key and "
     "compares it against the stored hash using a timing-safe comparison. Last-used "
     "timestamps are batched in memory and written to disk every 30 seconds to avoid "
     "a disk write on every API call."),

    # ------------------------------------------------------------------
    ("h1", "15.  Component Reference"),
    ("table", {
        "headers": ["Component", "Source File", "Responsibility"],
        "rows": [
            ["Scan Manager",     "app/scan_manager.py",        "Scan lifecycle, state, history, metrics"],
            ["Agent Graph",      "app/agent_graph.py",         "LangGraph state machine and all agent nodes"],
            ["Investigator",     "agents/investigator.py",     "Language model verdict on each finding"],
            ["Verifier",         "agents/verifier.py",         "PoV generation and validation orchestration"],
            ["Heuristic Scout",  "agents/heuristic_scout.py",  "Regex-based candidate discovery"],
            ["LLM Scout",        "agents/llm_scout.py",        "Language model candidate discovery"],
            ["Code Ingester",    "agents/ingest_codebase.py",  "ChromaDB embedding and retrieval"],
            ["Static Validator", "agents/static_validator.py", "Pattern-based PoV script verification"],
            ["Unit Test Runner", "agents/unit_test_runner.py", "Subprocess-based PoV execution"],
            ["Docker Runner",    "agents/docker_runner.py",    "Sandboxed container PoV execution"],
            ["Policy Agent",     "app/policy.py",              "Model routing decisions"],
            ["Learning Store",   "app/learning_store.py",      "SQLite investigation history"],
            ["Report Generator", "app/report_generator.py",    "JSON and PDF report output"],
            ["Git Handler",      "app/git_handler.py",         "Repository cloning and validation"],
            ["Source Handler",   "app/source_handler.py",      "ZIP extraction and code paste handling"],
            ["Webhook Handler",  "app/webhook_handler.py",     "GitHub and GitLab webhook processing"],
            ["Auth",             "app/auth.py",                "API key validation and rate limiting"],
            ["Config",           "app/config.py",              "All environment settings"],
            ["Prompts",          "prompts.py",                 "All language model prompt templates"],
        ],
        "col_widths_pdf":     [105, 165, 170],
        "col_widths_docx_cm": [4.0, 6.0, 6.0],
    }),
]


# ============================================================
# PDF  (fpdf2)
# ============================================================

from fpdf import FPDF
from pathlib import Path


FONT_DIR = "/usr/share/fonts/truetype/dejavu"

class DocPDF(FPDF):
    DARK   = (30,  30,  30)
    MID    = (80,  80,  80)
    LIGHT  = (200, 200, 200)
    XLIGHT = (240, 240, 240)
    WHITE  = (255, 255, 255)

    def __init__(self):
        super().__init__(orientation="P", unit="pt", format="A4")
        self.set_auto_page_break(auto=True, margin=60)
        self.set_margins(left=60, top=60, right=60)
        # Register DejaVu fonts for full Unicode support
        self.add_font("Body",     "",   f"{FONT_DIR}/DejaVuSans.ttf")
        self.add_font("Body",     "B",  f"{FONT_DIR}/DejaVuSans-Bold.ttf")
        self.add_font("Body",     "I",  f"{FONT_DIR}/DejaVuSans.ttf")
        self.add_font("Mono",     "",   f"{FONT_DIR}/DejaVuSansMono.ttf")
        self.add_font("Mono",     "B",  f"{FONT_DIR}/DejaVuSansMono-Bold.ttf")
        self.set_font("Body", size=10)
        self._header_drawn = False

    # ---- page header / footer ----
    def header(self):
        if not self._header_drawn:
            return
        self.set_y(16)
        self.set_font("Body", "I", 8)
        self.set_text_color(*self.MID)
        self.cell(0, 10, "AutoPoV -- Application Flow and System Reference", align="L")
        self.set_text_color(*self.DARK)

    def footer(self):
        self.set_y(-30)
        self.set_font("Body", "", 8)
        self.set_text_color(*self.MID)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")
        self.set_text_color(*self.DARK)

    # ---- helpers ----
    def usable_w(self):
        return self.w - self.l_margin - self.r_margin

    def gap(self, pt=6):
        self.ln(pt)

    # ---- cover page ----
    def cover(self, title, subtitle, sub2):
        self.add_page()
        self.set_y(self.h * 0.35)
        self.set_font("Body", "B", 28)
        self.set_text_color(*self.DARK)
        self.cell(0, 36, title, align="C", new_x="LMARGIN", new_y="NEXT")
        self.gap(8)
        self.set_font("Body", "", 14)
        self.set_text_color(*self.MID)
        self.cell(0, 20, subtitle, align="C", new_x="LMARGIN", new_y="NEXT")
        self.gap(4)
        self.set_font("Body", "I", 12)
        self.cell(0, 18, sub2, align="C", new_x="LMARGIN", new_y="NEXT")
        self.gap(20)
        self.set_draw_color(*self.LIGHT)
        self.set_line_width(0.8)
        self.line(self.l_margin, self.get_y(),
                  self.w - self.r_margin, self.get_y())
        self.set_text_color(*self.DARK)
        self._header_drawn = True

    # ---- section heading ----
    def h1(self, text):
        self.add_page()
        self.set_font("Body", "B", 14)
        self.set_text_color(*self.DARK)
        self.cell(0, 20, text, new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(*self.MID)
        self.set_line_width(0.6)
        self.line(self.l_margin, self.get_y(),
                  self.w - self.r_margin, self.get_y())
        self.gap(6)

    # ---- sub-heading ----
    def h2(self, text):
        self.gap(8)
        self.set_font("Body", "B", 11)
        self.set_text_color(*self.DARK)
        self.cell(0, 16, text, new_x="LMARGIN", new_y="NEXT")
        self.gap(3)

    # ---- body paragraph ----
    def body(self, text):
        self.set_font("Body", "", 10.5)
        self.set_text_color(*self.DARK)
        self.multi_cell(self.usable_w(), 15, text, new_x="LMARGIN", new_y="NEXT")
        self.gap(4)

    # ---- bullet list ----
    def bullets(self, items):
        self.set_font("Body", "", 10.5)
        self.set_text_color(*self.DARK)
        bullet = "    -  "
        indent = 14.0
        uw = self.usable_w() - indent
        for item in items:
            # Print bullet marker at left margin
            x_saved = self.get_x()
            self.set_x(self.l_margin)
            # Use multi_cell with first-line indent trick via manual positioning
            lines_text = bullet + item
            self.set_x(self.l_margin + indent)
            self.multi_cell(uw, 15, item, new_x="LMARGIN", new_y="NEXT")
            # go back and draw bullet on the first line of that cell
            # (already rendered above - just use prefix instead)
        # Re-render properly: write bullet then text on same line
        # Reset and do it properly with x positioning
        self.gap(4)

    # ---- REPLACEMENT bullet implementation ----
    def _bullets_impl(self, items):
        self.set_font("Body", "", 10.5)
        self.set_text_color(*self.DARK)
        import textwrap as _tw
        chars_per_line = max(20, int((self.usable_w() - 18) / (10.5 * 0.48)))
        for item in items:
            wrapped = _tw.wrap(item, chars_per_line) or [item]
            for li, wline in enumerate(wrapped):
                prefix = "  -  " if li == 0 else "      "
                self.set_x(self.l_margin)
                self.cell(self.usable_w(), 15, prefix + wline,
                          new_x="LMARGIN", new_y="NEXT")
        self.gap(4)

    # ---- code block ----
    def code_block(self, lines):
        self.gap(4)
        x0 = self.l_margin
        w  = self.usable_w()
        line_h = 13
        box_h  = len(lines) * line_h + 12
        if self.get_y() + box_h > self.h - self.b_margin:
            self.add_page()
        y0 = self.get_y()
        self.set_fill_color(*self.XLIGHT)
        self.set_draw_color(*self.LIGHT)
        self.set_line_width(0.4)
        self.rect(x0, y0, w, box_h, style="FD")
        self.set_xy(x0 + 8, y0 + 6)
        self.set_font("Mono", "", 9)
        self.set_text_color(*self.DARK)
        for line in lines:
            self.cell(w - 16, line_h, line if line else " ",
                      new_x="LMARGIN", new_y="NEXT")
            self.set_x(x0 + 8)
        self.set_y(y0 + box_h)
        self.gap(6)

    # ---- pipeline diagram ----
    def pipeline(self, lines):
        self.gap(4)
        x0   = self.l_margin
        w    = self.usable_w()
        lh   = 12
        box_h = len(lines) * lh + 14
        if self.get_y() + box_h > self.h - self.b_margin:
            self.add_page()
        y0 = self.get_y()
        self.set_fill_color(*self.XLIGHT)
        self.set_draw_color(*self.LIGHT)
        self.set_line_width(0.4)
        self.rect(x0, y0, w, box_h, style="FD")
        self.set_xy(x0 + 8, y0 + 7)
        self.set_font("Mono", "", 8.5)
        self.set_text_color(*self.DARK)
        for line in lines:
            self.cell(w - 16, lh, line if line else " ",
                      new_x="LMARGIN", new_y="NEXT")
            self.set_x(x0 + 8)
        self.set_y(y0 + box_h)
        self.gap(6)

    # ---- table ----
    def table(self, headers, rows, col_widths):
        self.gap(4)
        row_h  = 14
        pad    = 5
        total_w = sum(col_widths)
        if self.get_y() + row_h * 2 > self.h - self.b_margin:
            self.add_page()
        y = self.get_y()
        x0 = self.l_margin
        # header background
        self.set_fill_color(45, 45, 45)
        self.set_draw_color(*self.LIGHT)
        self.set_line_width(0.3)
        self.rect(x0, y, total_w, row_h + pad * 2, style="F")
        # header text
        self.set_font("Body", "B", 9)
        self.set_text_color(*self.WHITE)
        cx = x0
        for i, h in enumerate(headers):
            self.set_xy(cx + pad, y + pad + 1)
            self.cell(col_widths[i] - pad * 2, row_h, h.upper())
            cx += col_widths[i]
        y += row_h + pad * 2

        # data rows
        self.set_text_color(*self.DARK)
        for ri, row in enumerate(rows):
            # compute row height (tallest cell)
            cell_line_counts = []
            for i, cell in enumerate(row):
                cw = col_widths[i] if i < len(col_widths) else 80
                chars_per_line = max(8, int((cw - pad * 2) / (9 * 0.50)))
                import textwrap as tw
                wrapped = tw.wrap(str(cell), chars_per_line) or [""]
                cell_line_counts.append(len(wrapped))
            max_lines = max(cell_line_counts)
            rh = max_lines * row_h + pad * 2

            if y + rh > self.h - self.b_margin:
                self.add_page()
                y = self.get_y()
                # redraw header on continuation
                self.set_fill_color(45, 45, 45)
                self.rect(x0, y, total_w, row_h + pad * 2, style="F")
                self.set_font("Body", "B", 9)
                self.set_text_color(*self.WHITE)
                cx = x0
                for i, h in enumerate(headers):
                    self.set_xy(cx + pad, y + pad + 1)
                    self.cell(col_widths[i] - pad * 2, row_h, h.upper())
                    cx += col_widths[i]
                y += row_h + pad * 2
                self.set_text_color(*self.DARK)

            # alternating row bg
            if ri % 2 == 0:
                self.set_fill_color(*self.XLIGHT)
            else:
                self.set_fill_color(*self.WHITE)
            self.rect(x0, y, total_w, rh, style="F")

            # cell text
            self.set_font("Body", "", 9)
            cx = x0
            for i, cell in enumerate(row):
                cw = col_widths[i] if i < len(col_widths) else 80
                chars_per_line = max(8, int((cw - pad * 2) / (9 * 0.50)))
                import textwrap as tw
                wrapped = tw.wrap(str(cell), chars_per_line) or [""]
                ly = y + pad
                for wl in wrapped:
                    self.set_xy(cx + pad, ly)
                    self.cell(cw - pad * 2, row_h, wl)
                    ly += row_h
                cx += col_widths[i]

            # bottom border
            self.set_draw_color(*self.LIGHT)
            self.line(x0, y + rh, x0 + total_w, y + rh)
            y += rh

        self.set_y(y)
        self.gap(8)


def build_pdf(output_path: str):
    pdf = DocPDF()
    pdf.set_creator("AutoPoV _generate_docs.py")
    pdf.set_title("AutoPoV Application Flow")

    for kind, data in SECTIONS:
        if kind == "cover":
            pdf.cover(*data)
        elif kind == "h1":
            pdf.h1(data)
        elif kind == "h2":
            pdf.h2(data)
        elif kind == "body":
            pdf.body(data)
        elif kind == "bullets":
            pdf._bullets_impl(data)
        elif kind == "table":
            pdf.table(data["headers"], data["rows"], data["col_widths_pdf"])
        elif kind == "code_block":
            pdf.code_block(data)
        elif kind == "pipeline":
            pdf.pipeline(data)
        elif kind == "page_break":
            pdf.add_page()

    pdf.output(output_path)
    print(f"PDF written: {output_path}  ({pdf.page} pages)")


# ============================================================
# DOCX  (python-docx)
# ============================================================

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    sz)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def build_docx(output_path: str):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    DARK_COLOR   = RGBColor(30,  30,  30)
    MID_COLOR    = RGBColor(80,  80,  80)
    WHITE_HEX    = "FFFFFF"
    HEADER_HEX   = "1E1E1E"   # near-black header
    ODD_ROW_HEX  = "F5F5F5"

    def add_cover(title, subtitle, sub2):
        # blank line at top for spacing
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = DARK_COLOR

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(14)
        r2.font.color.rgb = MID_COLOR

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(sub2)
        r3.italic = True
        r3.font.size = Pt(12)
        r3.font.color.rgb = MID_COLOR
        doc.add_page_break()

    def add_h1(text):
        p = doc.add_heading(level=1)
        p.clear()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = DARK_COLOR
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)

    def add_h2(text):
        p = doc.add_heading(level=2)
        p.clear()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        run.font.color.rgb = DARK_COLOR
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)

    def add_body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size  = Pt(10.5)
            run.font.color.rgb = DARK_COLOR

    def add_bullets(items):
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.size = Pt(10.5)
            run.font.color.rgb = DARK_COLOR
        doc.add_paragraph("")

    def add_code_block(lines):
        p_before = doc.add_paragraph()
        p_before.paragraph_format.space_after = Pt(2)
        for line in lines:
            p = doc.add_paragraph()
            run = p.add_run(line if line else " ")
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_COLOR
            p.paragraph_format.left_indent   = Cm(0.8)
            p.paragraph_format.space_before  = Pt(0)
            p.paragraph_format.space_after   = Pt(0)
            # shading on paragraph via pPr
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "F0F0F0")
            pPr.append(shd)
        doc.add_paragraph("")

    def add_pipeline(lines):
        add_code_block(lines)

    def add_table(headers, rows, col_widths_cm):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        # header row
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h.upper()
            set_cell_bg(hdr_cells[i], HEADER_HEX)
            set_cell_borders(hdr_cells[i], "333333", "6")
            for run in hdr_cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # data rows
        for ri, row in enumerate(rows):
            cells = tbl.add_row().cells
            bg = ODD_ROW_HEX if ri % 2 == 0 else WHITE_HEX
            for i, val in enumerate(row):
                cells[i].text = str(val)
                set_cell_bg(cells[i], bg)
                set_cell_borders(cells[i], "CCCCCC", "4")
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = DARK_COLOR
        # column widths
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_cm):
                    cell.width = Cm(col_widths_cm[i])
        doc.add_paragraph("")

    # ---- render all sections ----
    for kind, data in SECTIONS:
        if kind == "cover":
            add_cover(*data)
        elif kind == "h1":
            add_h1(data)
        elif kind == "h2":
            add_h2(data)
        elif kind == "body":
            add_body(data)
        elif kind == "bullets":
            add_bullets(data)
        elif kind == "table":
            add_table(data["headers"], data["rows"], data["col_widths_docx_cm"])
        elif kind == "code_block":
            add_code_block(data)
        elif kind == "pipeline":
            add_pipeline(data)
        elif kind == "page_break":
            doc.add_page_break()

    doc.save(output_path)
    print(f"DOCX written: {output_path}")


# ============================================================
# Main
# ============================================================
if __name__ == "__main__":
    base = Path("/home/user/AutoPoV")
    build_pdf(str(base / "DOCS_APPLICATION_FLOW.pdf"))
    build_docx(str(base / "DOCS_APPLICATION_FLOW.docx"))
    print("Done.")
